from __future__ import annotations

import hashlib
import re
from pathlib import Path

from openpyxl import load_workbook

from ..catalogs import normalize_code_10
from .models import KnowledgeChunk, KnowledgeDocument

_SUPPORTED_EXTENSIONS = {
    ".txt": "text",
    ".md": "markdown",
    ".pdf": "pdf",
    ".docx": "docx",
    ".csv": "spreadsheet",
    ".xlsx": "spreadsheet",
    ".xls": "spreadsheet",
    ".json": "json",
}


def _collapse_spaces(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _extract_code_mentions(text: str) -> tuple[str, ...]:
    out: list[str] = []
    seen: set[str] = set()
    for match in re.findall(r"(?:\d[\d\s./-]{8,18}\d)", text or ""):
        code = normalize_code_10(match)
        if not code or code in seen:
            continue
        seen.add(code)
        out.append(code)
    return tuple(out)


def _split_text(text: str, *, chunk_size: int, chunk_overlap: int) -> list[str]:
    normalized_lines = [line.strip() for line in str(text or "").splitlines()]
    paragraphs = [line for line in normalized_lines if line]
    if not paragraphs:
        return []
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        candidate = f"{current}\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= chunk_size:
            current = candidate
            continue
        if current:
            chunks.append(current)
        if len(paragraph) <= chunk_size:
            current = paragraph
            continue
        start = 0
        step = max(1, chunk_size - max(0, chunk_overlap))
        while start < len(paragraph):
            piece = paragraph[start : start + chunk_size].strip()
            if piece:
                chunks.append(piece)
            start += step
        current = ""
    if current:
        chunks.append(current)
    return chunks


class KnowledgeChunkingService:
    def __init__(self, *, chunk_size: int = 1400, chunk_overlap: int = 220) -> None:
        self._chunk_size = max(400, int(chunk_size))
        self._chunk_overlap = max(0, min(int(chunk_overlap), self._chunk_size // 2))

    def scan_documents(self, root_dir: Path, *, source_kind: str) -> list[KnowledgeDocument]:
        root = Path(root_dir)
        if not root.exists():
            return []
        documents: list[KnowledgeDocument] = []
        for path in sorted(root.rglob("*")):
            if not path.is_file():
                continue
            document_type = _SUPPORTED_EXTENSIONS.get(path.suffix.lower())
            if not document_type:
                continue
            parser = self._parser_name_for_suffix(path.suffix.lower())
            stats = path.stat()
            documents.append(
                KnowledgeDocument(
                    source_path=str(path.resolve()),
                    relative_path=path.relative_to(root).as_posix(),
                    source_kind=source_kind,
                    document_type=document_type,
                    parser=parser,
                    file_sha256=self._sha256(path),
                    size_bytes=int(stats.st_size),
                    modified_at=float(stats.st_mtime),
                )
            )
        return documents

    def build_chunks(self, documents: list[KnowledgeDocument]) -> list[KnowledgeChunk]:
        chunks: list[KnowledgeChunk] = []
        for document in documents:
            source_path = Path(document.source_path)
            text = self.extract_text_from_file(source_path)
            parts = _split_text(
                text,
                chunk_size=self._chunk_size,
                chunk_overlap=self._chunk_overlap,
            )
            for chunk_index, part in enumerate(parts, start=1):
                clean_text = part.strip()
                if not clean_text:
                    continue
                section_context = clean_text.splitlines()[0][:160] if clean_text else ""
                chunk_id = f"{document.source_kind}:{document.relative_path}:{chunk_index}"
                chunks.append(
                    KnowledgeChunk(
                        chunk_id=chunk_id,
                        source_path=document.source_path,
                        relative_path=document.relative_path,
                        source_kind=document.source_kind,
                        document_type=document.document_type,
                        chunk_index=chunk_index,
                        text=clean_text,
                        section_context=section_context,
                        mentioned_codes=_extract_code_mentions(clean_text),
                        text_length=len(clean_text),
                    )
                )
        return chunks

    def extract_text_from_file(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md", ".json", ".csv"}:
            return self._read_text_file(path)
        if suffix == ".pdf":
            return self._read_pdf(path)
        if suffix == ".docx":
            return self._read_docx(path)
        if suffix in {".xlsx", ".xls"}:
            return self._read_spreadsheet(path)
        return ""

    @staticmethod
    def _parser_name_for_suffix(suffix: str) -> str:
        if suffix == ".pdf":
            return "pypdf2"
        if suffix == ".docx":
            return "python-docx"
        if suffix in {".xlsx", ".xls", ".csv"}:
            return "spreadsheet"
        return "text"

    @staticmethod
    def _sha256(path: Path) -> str:
        digest = hashlib.sha256()
        with path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()

    @staticmethod
    def _read_text_file(path: Path) -> str:
        for encoding in ("utf-8", "cp1251", "latin-1"):
            try:
                return path.read_text(encoding=encoding)
            except Exception:
                continue
        return ""

    @staticmethod
    def _read_pdf(path: Path) -> str:
        try:
            from PyPDF2 import PdfReader
        except Exception:
            return ""
        try:
            reader = PdfReader(str(path))
        except Exception:
            return ""
        pages: list[str] = []
        for page in reader.pages:
            try:
                pages.append(str(page.extract_text() or ""))
            except Exception:
                continue
        return "\n".join(item for item in pages if item)

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            from docx import Document
        except Exception:
            return ""
        try:
            document = Document(str(path))
        except Exception:
            return ""
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = _collapse_spaces(paragraph.text)
            if text:
                blocks.append(text)
        for table in document.tables:
            for row in table.rows:
                values = [_collapse_spaces(cell.text) for cell in row.cells]
                row_text = " | ".join(value for value in values if value)
                if row_text:
                    blocks.append(row_text)
        return "\n".join(blocks)

    @staticmethod
    def _read_spreadsheet(path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".csv":
            return KnowledgeChunkingService._read_text_file(path)
        if suffix == ".xls":
            try:
                import xlrd  # type: ignore
            except Exception:
                return ""
            try:
                workbook = xlrd.open_workbook(str(path))
            except Exception:
                return ""
            rows: list[str] = []
            for sheet in workbook.sheets():
                rows.append(f"[sheet] {sheet.name}")
                for row_idx in range(sheet.nrows):
                    values = [_collapse_spaces(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
                    line = " | ".join(value for value in values if value)
                    if line:
                        rows.append(line)
            return "\n".join(rows)
        try:
            workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
        except Exception:
            return ""
        rows: list[str] = []
        try:
            for sheet in workbook.worksheets:
                rows.append(f"[sheet] {sheet.title}")
                for raw_row in sheet.iter_rows(values_only=True):
                    values = [_collapse_spaces(item) for item in raw_row]
                    line = " | ".join(value for value in values if value)
                    if line:
                        rows.append(line)
        finally:
            workbook.close()
        return "\n".join(rows)
